from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x0F, 0x1B, 0x2D)
GOLD  = RGBColor(0xC8, 0x95, 0x2E)
CREAM = RGBColor(0xF6, 0xF1, 0xE9)
MUTED = RGBColor(0x5B, 0x6E, 0x8A)
BODY  = RGBColor(0x1A, 0x29, 0x42)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_para_shading(para, color_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def set_para_border_left(para, color_hex, size=36):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)

def add_run(para, text, bold=False, italic=False, size=12, color=None, font='Inter'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, size=22, color=None, bold=True, font='Montserrat', space_before=18, space_after=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    add_run(p, text, bold=bold, size=size, color=color or NAVY, font=font)
    return p

def body(doc, text, size=11.5, color=None, italic=False, bold=False, space_before=6, space_after=8, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        pf.left_indent = Inches(indent)
    add_run(p, text, italic=italic, bold=bold, size=size, color=color or BODY)
    return p

def gold_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C8952E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def spacer(doc, size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(size)
    return p

def step_number(doc, n):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(0)
    add_run(p, f'{n:02d}', bold=True, size=48, color=GOLD, font='Montserrat')
    return p

def pull_quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(16)
    pf.left_indent = Inches(0.3)
    pf.right_indent = Inches(0.3)
    set_para_shading(p, '0F1B2D')
    add_run(p, f'"{text}"', italic=True, size=12, color=CREAM)
    return p

def callout_line(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.left_indent = Inches(0.15)
    set_para_shading(p, 'F6F1E9')
    set_para_border_left(p, 'C8952E', size=36)
    add_run(p, text, bold=True, size=11.5, color=NAVY)
    return p

def action_box_header(doc, label='SUA AÇÃO PARA ESTE PASSO'):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(6)
    pf.left_indent = Inches(0.2)
    set_para_shading(p, 'F6F1E9')
    add_run(p, label, bold=True, size=9, color=GOLD, font='Inter')
    return p

def action_item(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.35)
    set_para_shading(p, 'F6F1E9')
    add_run(p, f'□  {text}', size=11, color=BODY)
    return p

def compliance(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(6)
    add_run(p, text, italic=True, size=9, color=MUTED)
    return p

def transition_line(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    add_run(p, text, italic=True, size=10.5, color=MUTED)
    return p

def section_head(doc, text, size=13):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(20)
    pf.space_after = Pt(8)
    add_run(p, text, bold=True, size=size, color=NAVY)
    return p

def bullet_item(doc, label, text, indent=0.2):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(5)
    pf.space_after = Pt(5)
    pf.left_indent = Inches(indent)
    if label:
        add_run(p, f'{label}  ', bold=True, size=11, color=NAVY)
        add_run(p, text, size=11, color=BODY)
    else:
        add_run(p, f'•  {text}', size=11, color=BODY)
    return p

def set_doc_margins(doc, top=1.0, bottom=0.9, left=1.1, right=1.1):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

# ── BUILD ────────────────────────────────────────────────────

doc = Document()
set_doc_margins(doc)

# ── CAPA ─────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.space_after = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_shading(p, '0F1B2D')
add_run(p, 'PROSPER IN AMERICA', bold=True, size=32, color=GOLD, font='Montserrat')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(6)
set_para_shading(p2, '0F1B2D')
add_run(p2, 'O Guia Financeiro Inicial do Imigrante', size=16, color=CREAM, font='Inter')

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(4)
p3.paragraph_format.space_after = Pt(4)
set_para_shading(p3, '0F1B2D')
add_run(p3, '7 Passos Para Construir Sua Vida nos Estados Unidos', italic=True, size=13, color=CREAM, font='Inter')

for _ in range(6):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(0)
    set_para_shading(sp, '0F1B2D')

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(0)
p4.paragraph_format.space_after = Pt(80)
set_para_shading(p4, '0F1B2D')
add_run(p4, 'prosperinamerica.com', size=10, color=MUTED, font='Inter')

doc.add_page_break()

# ── CARTA DA WELLEN ──────────────────────────────────────────
heading(doc, 'Uma Nota da Wellen', size=20, space_before=12)
gold_rule(doc)
spacer(doc, 14)

for line in [
    "Você não veio pra cá pra ficar no mesmo lugar.",
    "",
    "Eu também não.",
    "",
    "Quando eu cheguei aqui, eu não sabia como abrir uma conta.",
    "Não entendia crédito.",
    "Não sabia nem por onde começar.",
    "",
    "Aprendi errando.",
    "",
    "Eu sei o que é trabalhar duro —",
    "e sentir que não está saindo do lugar.",
    "",
    "Esse guia existe por um motivo:",
    "",
    "Te dar clareza.",
    "Te dar direção.",
    "E te ajudar a não pagar o preço dos erros que ninguém te explicou.",
    "",
    "Leia isso com atenção.",
    "",
    "E mais importante:",
    "coloque em prática.",
    "",
    "Se quiser ajuda pra aplicar isso na sua realidade,",
    "fala com a gente direto no WhatsApp.",
]:
    if line == "":
        spacer(doc, 6)
    else:
        body(doc, line, space_before=2, space_after=3)

spacer(doc, 14)
p_sig = doc.add_paragraph()
p_sig.paragraph_format.space_before = Pt(10)
p_sig.paragraph_format.space_after = Pt(2)
add_run(p_sig, '— Wellen', bold=True, size=12, color=NAVY)
p_brand = doc.add_paragraph()
p_brand.paragraph_format.space_before = Pt(0)
p_brand.paragraph_format.space_after = Pt(6)
add_run(p_brand, 'Prosper In America', size=11, color=GOLD)

doc.add_page_break()

# ── COMO USAR ────────────────────────────────────────────────
heading(doc, 'Como Usar Este Guia', size=20, space_before=12)
gold_rule(doc)
spacer(doc, 10)

body(doc, 'Este guia tem 7 passos.')
body(doc, 'Eles seguem uma sequência — mas você não precisa começar do início.')
spacer(doc, 8)
body(doc, 'Se você já tem conta no banco, pule para o Passo 3.')
body(doc, 'Se crédito é o seu maior gap, comece pelo Passo 4.')
body(doc, 'Se você está preocupado com a proteção da sua família, vá direto para o Passo 5.')
spacer(doc, 10)
body(doc, 'Encontre onde você está. Comece por aí.', bold=True, color=NAVY)

spacer(doc, 16)
section_head(doc, 'Passos deste guia')

for num, title in [
    ('01', 'Por Que a Maioria dos Imigrantes Fica Parada no Lugar'),
    ('02', 'Coloque Seus Números em Ordem'),
    ('03', 'Abra a Conta Certa'),
    ('04', 'Construa Seu Crédito'),
    ('05', 'Proteja o Que Você Está Construindo'),
    ('06', 'Aumente Sua Renda'),
    ('07', 'Construa Algo que Dure'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.2)
    set_para_shading(p, 'F6F1E9')
    add_run(p, f'{num}  ', bold=True, size=12, color=GOLD, font='Montserrat')
    add_run(p, title, size=12, color=NAVY, font='Inter')

doc.add_page_break()

# ── PASSO 1 ──────────────────────────────────────────────────
step_number(doc, 1)
heading(doc, 'Por Que a Maioria dos Imigrantes Fica Parada no Lugar', size=22, space_before=4)
body(doc, 'Não é falta de esforço. É falta de estrutura.', italic=True, color=MUTED, space_before=0, space_after=12)
gold_rule(doc)
spacer(doc, 10)

body(doc, 'Você trabalha muito.')
body(doc, 'Isso nunca foi o problema.')
spacer(doc, 12)
body(doc, 'O problema é não saber o que fazer com o que você ganha.')
spacer(doc, 12)
body(doc, 'Você veio pra construir algo.')
body(doc, 'Mas os anos passam — e nada muda.')
spacer(doc, 10)
body(doc, 'O dinheiro entra.')
body(doc, 'O dinheiro sai.')
body(doc, 'Nada fica.')
spacer(doc, 12)
body(doc, 'Isso não é falta de esforço.', bold=True, color=NAVY)
spacer(doc, 8)
body(doc, 'É falta de direção.', bold=True, color=NAVY)
spacer(doc, 16)

section_head(doc, 'Ninguém te deu o mapa.')
body(doc, 'O sistema financeiro aqui não se explica.')
body(doc, 'Ele espera que você já saiba tudo.')
spacer(doc, 10)
body(doc, 'Se você não sabia — você foi aprendendo no escuro.')
spacer(doc, 14)

section_head(doc, 'Você decidiu baseado no que apareceu:')
for line in [
    'o que alguém falou,',
    'o que viu no WhatsApp,',
    'o que parecia fazer sentido na hora.',
]:
    body(doc, line, indent=0.25, space_before=4, space_after=4)

spacer(doc, 12)
body(doc, 'E algumas dessas decisões já estão te custando.', bold=True, color=NAVY)
spacer(doc, 8)
body(doc, 'Você só ainda não percebeu o quanto.', bold=True, color=NAVY)
spacer(doc, 16)

section_head(doc, 'O que muda quando você tem a informação certa')
for line in [
    'Você para de tomar decisões com base no medo ou em suposições.',
    'Você sabe qual passo vem primeiro — e por quê.',
    'Você para de perder tempo com as coisas erradas.',
    'O caminho fica claro.',
]:
    body(doc, line, space_before=4, space_after=4)

spacer(doc, 12)
pull_quote(doc, 'A diferença quase nunca é esforço. Quase sempre é informação — e saber em qual ordem agir.')
spacer(doc, 14)

action_box_header(doc)
p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Pt(4)
p_sub.paragraph_format.space_after = Pt(6)
p_sub.paragraph_format.left_indent = Inches(0.2)
set_para_shading(p_sub, 'F6F1E9')
add_run(p_sub, 'Identifique a área em que você se sente mais atrasado:', size=11, color=BODY)

for item in ['Conta bancária', 'Histórico de crédito', 'Declaração de imposto', 'Proteger minha família', 'Construir renda']:
    action_item(doc, item)

spacer(doc, 8)
transition_line(doc, 'Quando você entende onde está, o próximo passo é descobrir o que você realmente tem acesso.')

doc.add_page_break()

# ── PASSO 2 ──────────────────────────────────────────────────
step_number(doc, 2)
heading(doc, 'Coloque Seus Números em Ordem', size=22, space_before=4)
body(doc, 'ITIN e SSN — o que cada um é e o que cada um abre.', italic=True, color=MUTED, space_before=0, space_after=12)
gold_rule(doc)
spacer(doc, 10)

body(doc, 'Você chega sabendo que precisa de algo pra começar.')
body(doc, 'Mas sem saber exatamente o que é — ou o que ele abre.')
spacer(doc, 16)

section_head(doc, 'SSN — Social Security Number')
body(doc, 'Emitido pela Social Security Administration para pessoas autorizadas a trabalhar nos EUA.')
spacer(doc, 6)
body(doc, 'Se você tem autorização de trabalho — por visto, green card ou outro status — você pode ser elegível para solicitá-lo.')
spacer(doc, 6)
body(doc, 'Se você não tem autorização de trabalho, não vai conseguir um SSN.')
body(doc, 'É aí que entra o ITIN.')
spacer(doc, 14)

section_head(doc, 'ITIN — Individual Taxpayer Identification Number')
body(doc, 'Emitido pelo IRS.')
body(doc, 'É um número de processamento fiscal — não é prova de status imigratório, não é autorização de trabalho, e não é um caminho para benefícios por si só.')
spacer(doc, 6)
body(doc, 'Existe para pessoas que têm obrigação de declarar imposto nos EUA mas não são elegíveis para o SSN.')
spacer(doc, 12)

section_head(doc, 'Para que o ITIN é comumente usado:', size=12)
for item in [
    'Abrir certas contas bancárias',
    'Declarar imposto de renda',
    'Construir histórico de crédito em algumas instituições',
    'Solicitar certos produtos financeiros',
]:
    bullet_item(doc, None, item)

spacer(doc, 12)
section_head(doc, 'Como funciona o pedido, de forma geral:', size=12)
body(doc, 'Você preenche o formulário W-7 do IRS.')
body(doc, 'Envia junto com sua declaração de imposto e documentos de identidade.')
spacer(doc, 6)
body(doc, 'Agentes certificados — chamados Certifying Acceptance Agents (CAAs) — podem ajudar no processo.')
body(doc, 'Muitas empresas contábeis e organizações sem fins lucrativos oferecem esse serviço.')
spacer(doc, 6)
body(doc, 'O tempo de processamento varia. Consulte o site do IRS para estimativas atuais.')
spacer(doc, 14)

section_head(doc, 'O que isso significa pra você')
body(doc, 'Se você tem SSN — você já tem uma vantagem.')
body(doc, 'Certifique-se de saber o que ele te dá e não te dá acesso.')
spacer(doc, 8)
body(doc, 'Se você tem ITIN — você tem um ponto de partida real.')
body(doc, 'Muitos bancos, credores e serviços aceitam.')
spacer(doc, 8)
body(doc, 'Se você não tem nenhum dos dois — descubra se o ITIN se aplica à sua situação.')
body(doc, 'É o ponto de partida para muitas pessoas sem autorização de trabalho.')
spacer(doc, 10)

callout_line(doc, 'Não espere ter o SSN pra começar a construir. Para muitas pessoas nessa situação, o ITIN é o que abre as primeiras portas de verdade.')
spacer(doc, 12)

action_box_header(doc)
for item in [
    'Confirme se você tem SSN ou ITIN',
    'Se não tiver nenhum, pesquise o processo do ITIN — um contador ou CAA pode te orientar',
    'Anote para que você quer usar o número (banco, crédito, declaração)',
]:
    action_item(doc, item)

spacer(doc, 8)
compliance(doc, 'Observação: A elegibilidade para o ITIN e SSN depende do status imigratório individual. Informação geral apenas — consulte um profissional especializado em tributos ou imigração.')
transition_line(doc, 'Com o número certo em mãos, o próximo passo é onde seu dinheiro vai morar.')

doc.add_page_break()

# ── PASSO 3 ──────────────────────────────────────────────────
step_number(doc, 3)
heading(doc, 'Abra a Conta Certa', size=22, space_before=4)
gold_rule(doc)
spacer(doc, 12)

p_strong3 = doc.add_paragraph()
p_strong3.paragraph_format.space_before = Pt(4)
p_strong3.paragraph_format.space_after = Pt(10)
add_run(p_strong3, 'Se a sua conta está errada, tudo o que vem depois fica mais difícil.', bold=True, size=14, color=NAVY)

spacer(doc, 6)
body(doc, 'Você não consegue construir crédito nos EUA sem uma conta.')
body(doc, 'Não dá pra pagar aluguel, contas e seguro de forma confiável sem uma.')
spacer(doc, 6)
body(doc, 'E se você abriu a conta errada — ou pulou esse passo — está pagando um preço que talvez ainda não tenha percebido.')
spacer(doc, 14)

section_head(doc, 'Por que a conta errada atrasa tudo')
body(doc, 'Nem toda conta bancária é igual.')
body(doc, 'Algumas foram feitas pra quem já está estabelecido.')
body(doc, 'Elas cobram tarifas mensais, exigem saldo mínimo e têm multas que vão drenando o que você está tentando construir.')
spacer(doc, 8)
body(doc, 'Para um imigrante começando do zero, uma conta com tarifa mensal de $15 e saldo mínimo de $500 não é só inconveniente.')
spacer(doc, 4)
callout_line(doc, 'Isso não é tarifa. Isso é um vazamento.')
spacer(doc, 10)

section_head(doc, 'O que procurar')
for label, text in [
    ('Aceita seus documentos:', 'Muitos bancos aceitam ITIN, passaporte e comprovante de endereço nos EUA — sem SSN obrigatório. Confirme diretamente antes de solicitar.'),
    ('Sem tarifas mensais:', 'Algumas contas cobram só por existir. Procure uma sem tarifa, ou que seja facilmente dispensada.'),
    ('Sem saldo mínimo:', 'Se você está começando do zero, não deveria ser penalizado por ter saldo baixo.'),
    ('Acesso online e pelo celular:', 'Básico. Se a conta não oferece, continue procurando.'),
    ('Caminho para crédito:', 'Alguns bancos oferecem produtos de construção de crédito para clientes existentes. Isso importa mais do que a maioria das pessoas percebe.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 12)
section_head(doc, 'O que evitar')
for label, text in [
    ('Cartão pré-pago como substituto de conta:', 'Não reporta ao bureau de crédito. Não constrói histórico bancário. É um recurso emergencial, não uma base.'),
    ('Contas com tarifas empilhadas:', 'Tarifa mensal + taxa de cheque especial + ATM fora da rede. Leia o resumo de tarifas antes de abrir qualquer conta.'),
    ('Esperar até ter SSN:', 'Em muitos casos, ITIN e passaporte são suficientes. Esperar custa tempo que você não precisa perder.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 12)
section_head(doc, 'Documentos normalmente solicitados', size=12)
for item in [
    'Passaporte (válido)',
    'ITIN ou SSN (quando exigido)',
    'Comprovante de endereço nos EUA — contrato de aluguel, conta de luz ou correspondência oficial',
]:
    bullet_item(doc, None, item)

spacer(doc, 8)
body(doc, 'Sempre ligue ou consulte o site da instituição antes de ir.', italic=True, color=MUTED)
body(doc, 'As exigências podem variar entre agências do mesmo banco.', italic=True, color=MUTED)
spacer(doc, 12)

pull_quote(doc, 'A conta certa não impressiona ninguém. Ela simplesmente funciona — sem tirar dinheiro enquanto faz isso.')
spacer(doc, 10)

p_warn3 = doc.add_paragraph()
p_warn3.paragraph_format.space_before = Pt(4)
p_warn3.paragraph_format.space_after = Pt(10)
add_run(p_warn3, 'Se esse passo estiver errado, tudo o que vem depois fica mais lento, mais caro e mais difícil de corrigir.', bold=True, size=11.5, color=NAVY)

action_box_header(doc)
for item in [
    'Pesquise 2–3 opções de conta que aceitam seus documentos atuais',
    'Compare as tarifas — procure o resumo completo, não só os benefícios anunciados',
    'Confirme os requisitos diretamente com a instituição antes de solicitar',
    'Abra uma conta que não te cobre nada pra existir',
]:
    action_item(doc, item)

spacer(doc, 8)
compliance(doc, 'Observação: A elegibilidade para abertura de conta varia por instituição e estado. Confirme os requisitos diretamente com o banco ou cooperativa de crédito antes de solicitar.')
transition_line(doc, 'Com a conta certa aberta, o sistema começa a olhar para o seu histórico de crédito.')

doc.add_page_break()

# ── PASSO 4 ──────────────────────────────────────────────────
step_number(doc, 4)
heading(doc, 'Construa Seu Crédito', size=22, space_before=4)
gold_rule(doc)
spacer(doc, 10)

p_strong4 = doc.add_paragraph()
p_strong4.paragraph_format.space_before = Pt(4)
p_strong4.paragraph_format.space_after = Pt(6)
add_run(p_strong4, 'Todo o seu histórico financeiro no Brasil não vale nada aqui.', bold=True, size=14, color=NAVY)
p_strong4b = doc.add_paragraph()
p_strong4b.paragraph_format.space_before = Pt(0)
p_strong4b.paragraph_format.space_after = Pt(12)
add_run(p_strong4b, 'Você começa do zero.', bold=True, size=14, color=NAVY)

body(doc, 'Muitos imigrantes chegam sem nenhum histórico de crédito nos EUA.')
body(doc, 'Não crédito ruim. Nenhum crédito.')
body(doc, 'Aos olhos do sistema financeiro, você ainda não existe.')
spacer(doc, 8)
body(doc, 'É por isso que as pessoas ficam presas por mais tempo do que esperavam.')
spacer(doc, 8)
body(doc, 'Você não pode apressar o histórico de crédito.')
body(doc, 'O que dá pra fazer é evitar os erros que te obrigam a recomeçar — e dar os passos certos cedo o suficiente pra que o tempo trabalhe a seu favor.')
spacer(doc, 14)

section_head(doc, 'O que realmente importa — os três fatores')
for label, text in [
    ('Histórico de pagamentos:', 'O fator mais importante. Pague em dia, sempre. Um atraso pode ficar no seu relatório por anos.'),
    ('Utilização do crédito:', 'Quanto da sua linha disponível você está usando. Alta utilização sinaliza pressão financeira. Mantenha os saldos baixos.'),
    ('Tempo de histórico:', 'Contas mais antigas ajudam. Fechar sua primeira conta pra fazer upgrade costuma ser um erro.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 8)
body(doc, 'Todo o resto — consultas de crédito, mix de produtos — importa menos. Foque nesses três primeiro.', italic=True, color=MUTED)
spacer(doc, 14)

section_head(doc, 'Como imigrantes costumam começar a construir crédito')
for label, text in [
    ('Cartão de crédito garantido (secured):', 'Um depósito vira seu limite. Use, pague em dia. O comportamento de pagamento constrói o histórico — não o depósito.'),
    ('Virar usuário autorizado:', 'Alguém de confiança te adiciona no cartão dele. O histórico deles pode aparecer no seu relatório — uma das formas mais rápidas de estabelecer uma pontuação inicial.'),
    ('Empréstimo para construção de crédito:', 'Oferecido por alguns bancos e cooperativas. Pagamentos mensais fixos reportados ao bureau. Pesquise disponibilidade onde você tem conta.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 14)
section_head(doc, 'O que evitar')
for label, text in [
    ('Atrasar um pagamento — mesmo uma vez:', 'Configure pagamento mínimo automático. Um atraso pode ficar no seu histórico por até sete anos.'),
    ('Maximizar o cartão:', 'Alta utilização prejudica sua pontuação mesmo que você pague depois. Mantenha os saldos bem abaixo do seu limite.'),
    ('Solicitar vários produtos de uma vez:', 'Cada pedido gera uma consulta. Pesquise primeiro, solicite uma vez.'),
    ('Achar que seu histórico financeiro do Brasil vale aqui:', 'Não vale. O sistema americano não tem acesso a ele. Isso não é punição — é como o sistema funciona.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 8)
callout_line(doc, 'Quem descobre isso tarde, olha pra trás e vê anos que não consegue recuperar.')
spacer(doc, 12)
pull_quote(doc, 'Você não pode apressar o histórico de crédito. O que você pode fazer é evitar os erros que te obrigam a recomeçar.')
spacer(doc, 12)

action_box_header(doc)
for item in [
    'Verifique se você tem histórico de crédito nos EUA em annualcreditreport.com (gratuito)',
    'Pesquise um cartão garantido ou produto de construção de crédito adequado à sua situação',
    'Configure pagamento automático em todas as contas existentes',
    'Identifique alguém de confiança que possa te adicionar como usuário autorizado — e tenha essa conversa',
]:
    action_item(doc, item)

spacer(doc, 8)
compliance(doc, 'Observação: Pontuação de crédito e decisões de crédito dependem de muitos fatores individuais. Informação educacional apenas — não é aconselhamento financeiro. Consulte um profissional financeiro qualificado.')
transition_line(doc, 'Enquanto seu histórico financeiro começa a se construir, a próxima pergunta é: o que acontece se tudo parar?')

doc.add_page_break()

# ── PASSO 5 ──────────────────────────────────────────────────
step_number(doc, 5)
heading(doc, 'Proteja o Que Você Está Construindo', size=22, space_before=4)
body(doc, 'Muitos imigrantes passam anos construindo — e não protegem nada disso.', bold=True, size=13, color=NAVY, space_before=0, space_after=12)
gold_rule(doc)
spacer(doc, 10)

body(doc, 'Esse é o passo que as pessoas mais deixam pra depois.')
spacer(doc, 8)
body(doc, 'Não porque não se importam com a família.')
body(doc, 'Porque estão focadas em sobreviver agora — e a proteção parece algo pra resolver depois, quando as coisas estiverem mais estáveis.')
spacer(doc, 8)
body(doc, 'Essa lógica é comum.')
body(doc, 'E é exatamente assim que famílias acabam em crise.')
spacer(doc, 14)

section_head(doc, 'Como o risco realmente parece')
p_scenario = doc.add_paragraph()
p_scenario.paragraph_format.space_before = Pt(8)
p_scenario.paragraph_format.space_after = Pt(6)
p_scenario.paragraph_format.left_indent = Inches(0.2)
p_scenario.paragraph_format.right_indent = Inches(0.2)
set_para_shading(p_scenario, 'F6F1E9')
add_run(p_scenario,
    'Você é quem sustenta a casa.\n\n'
    'Você trabalha.\n'
    'Paga o aluguel.\n'
    'Manda dinheiro pro Brasil.\n\n'
    'Todo mês, é você.\n\n'
    'Agora pensa comigo:\n\n'
    'Se você parar por 6 meses —\n'
    'quem segura tudo?\n\n'
    'Se algo acontecer com você —\n'
    'quem sustenta sua família?\n\n'
    'E a família no Brasil que depende de você?',
    size=11.5, color=BODY)

spacer(doc, 14)
body(doc, 'Isso não é sobre medo.', bold=True, color=NAVY)
spacer(doc, 8)
body(doc, 'É sobre responsabilidade.', bold=True, color=NAVY)
spacer(doc, 8)
body(doc, 'E responsabilidade exige preparação.', bold=True, color=NAVY)
spacer(doc, 10)

section_head(doc, 'O que proteção realmente significa')
for label, text in [
    ('Seguro de vida:', 'Uma apólice paga um valor determinado à sua família se você falecer. É uma estrutura simples: você paga uma mensalidade, e se o pior acontecer, sua família recebe um valor que pode cobrir despesas finais, renda perdida e tempo pra se reorganizar.'),
    ('Plano de saúde:', 'Dívida médica é uma das principais causas de colapso financeiro nos EUA — mesmo para famílias que trabalham. Uma única internação sem cobertura pode apagar anos de poupança.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 14)
section_head(doc, 'Seguro de vida e status imigratório')
body(doc, 'Muitas famílias assumem que precisam de SSN pra se qualificar para seguro de vida.')
body(doc, 'Essa suposição custa anos de cobertura que elas poderiam ter tido.')
spacer(doc, 8)
body(doc, 'Algumas apólices aceitam solicitantes com ITIN.')
body(doc, 'O processo de análise varia por apólice e operadora.')
body(doc, 'Mas a opção existe.')
spacer(doc, 8)
callout_line(doc, 'Se você tem esperado até ter seu SSN, descubra o que você já pode ser elegível agora.')
spacer(doc, 12)

section_head(doc, 'O que evitar')
for label, text in [
    ('Esperar até se sentir estável:', 'Não existe estável. Estabilidade se constrói, não se espera. Quem espera a hora certa costuma esperar demais.'),
    ('Achar que não vai se qualificar:', 'Muita gente nunca verifica. Essa suposição está errada com mais frequência do que as pessoas imaginam.'),
    ('Escolher cobertura só pelo preço:', 'A opção mais barata e a opção certa nem sempre são a mesma coisa.'),
    ('Deixar essa conversa pra quando algo forçar:', 'A hora de pensar em proteger sua família não é durante uma crise. É antes.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 12)
pull_quote(doc, 'Proteção não é o que você constrói depois que está estável. Faz parte do que torna a estabilidade possível.')
spacer(doc, 10)

p_delay = doc.add_paragraph()
p_delay.paragraph_format.space_before = Pt(4)
p_delay.paragraph_format.space_after = Pt(6)
add_run(p_delay, 'Esse passo quase sempre é adiado até que algo force a decisão.', bold=True, size=11.5, color=NAVY)
p_delay2 = doc.add_paragraph()
p_delay2.paragraph_format.space_before = Pt(0)
p_delay2.paragraph_format.space_after = Pt(12)
add_run(p_delay2, 'Quando isso acontece, as opções são menores.', bold=True, size=11.5, color=NAVY)

action_box_header(doc)
for item in [
    'Pergunte a si mesmo: se eu não pudesse trabalhar por 6 meses, o que aconteceria com minha família?',
    'Pesquise opções de plano de saúde — comece pelo healthcare.gov ou um assistente local',
    'Verifique se seus documentos atuais permitem solicitar seguro de vida',
    'Converse com um profissional licenciado — não pra comprar nada, mas pra entender o que existe',
]:
    action_item(doc, item)

spacer(doc, 8)
compliance(doc, 'Observação: A elegibilidade para seguros, coberturas e programas de saúde varia por circunstâncias, status imigratório e estado. Conteúdo informacional apenas — não é aconselhamento de seguros. Consulte um profissional de seguros ou saúde licenciado.')

# ── PÁGINA CTA ───────────────────────────────────────────────
doc.add_page_break()

spacer(doc, 50)
heading(doc, 'Precisa de ajuda pra aplicar\nisso na sua situação?', size=22, space_before=12, space_after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
gold_rule(doc)
spacer(doc, 20)

body(doc, 'Você não precisa fazer isso sozinho.', bold=True, size=13, color=NAVY, space_before=6, space_after=12)

p_cta1 = doc.add_paragraph()
p_cta1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cta1.paragraph_format.space_before = Pt(6)
p_cta1.paragraph_format.space_after = Pt(4)
add_run(p_cta1, 'Se quiser clareza sobre o seu caso específico,', size=12, color=BODY)

p_cta2 = doc.add_paragraph()
p_cta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cta2.paragraph_format.space_before = Pt(0)
p_cta2.paragraph_format.space_after = Pt(6)
add_run(p_cta2, 'fale com a gente direto no WhatsApp.', size=12, color=BODY)

spacer(doc, 32)

p_btn = doc.add_paragraph()
p_btn.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_btn.paragraph_format.space_before = Pt(10)
p_btn.paragraph_format.space_after = Pt(10)
set_para_shading(p_btn, 'C8952E')
add_run(p_btn, 'Falar no WhatsApp', bold=True, size=14, color=WHITE)

spacer(doc, 10)
p_url = doc.add_paragraph()
p_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_url.paragraph_format.space_before = Pt(4)
p_url.paragraph_format.space_after = Pt(4)
add_run(p_url, 'wa.me/13526303930', italic=True, size=10, color=MUTED)

# ── SALVAR ───────────────────────────────────────────────────
out_path = '/Users/miriampalma/AI-OS/projects/prosper-landing/public/assets/prosper_in_america_guia_pt.docx'
doc.save(out_path)
print(f'Salvo: {out_path}')
