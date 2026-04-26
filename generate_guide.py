from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color palette ────────────────────────────────────────────
NAVY    = RGBColor(0x0F, 0x1B, 0x2D)
GOLD    = RGBColor(0xC8, 0x95, 0x2E)
CREAM   = RGBColor(0xF6, 0xF1, 0xE9)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
MUTED   = RGBColor(0x5B, 0x6E, 0x8A)
BODY    = RGBColor(0x1A, 0x29, 0x42)

# ── Helpers ──────────────────────────────────────────────────

def set_page_background(doc, color_hex):
    """Set document page background color."""
    background = OxmlElement('w:background')
    background.set(qn('w:color'), color_hex)
    background.set(qn('w:themeColor'), 'none')
    doc.element.insert(0, background)
    settings = doc.settings.element
    ds = OxmlElement('w:displayBackgroundShape')
    settings.insert(0, ds)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_para_shading(para, color_hex):
    """Set paragraph background shading."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def set_para_border_left(para, color_hex, size=36):
    """Add left border to paragraph (gold rule effect)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)

def add_run(para, text, bold=False, italic=False, size=12, color=None, font='Inter'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, size=22, color=None, bold=True, font='Montserrat', space_before=18, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    add_run(p, text, bold=bold, size=size, color=color or NAVY, font=font)
    return p

def body(doc, text, size=11.5, color=None, italic=False, bold=False, space_before=4, space_after=6, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        pf.left_indent = Inches(indent)
    add_run(p, text, italic=italic, bold=bold, size=size, color=color or BODY)
    return p

def gold_rule(doc):
    """Thin gold horizontal rule via paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C8952E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def spacer(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(size)
    return p

def step_number(doc, n):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(0)
    add_run(p, f'{n:02d}', bold=True, size=48, color=GOLD, font='Montserrat')
    return p

def pull_quote(doc, text):
    """Navy background pull quote block."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.left_indent = Inches(0.3)
    pf.right_indent = Inches(0.3)
    set_para_shading(p, '0F1B2D')
    add_run(p, f'“{text}”', italic=True, size=12, color=CREAM)
    return p

def callout_line(doc, text, bold=True):
    """Gold left-border callout with cream background."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)
    pf.left_indent = Inches(0.15)
    set_para_shading(p, 'F6F1E9')
    set_para_border_left(p, 'C8952E', size=36)
    add_run(p, text, bold=bold, size=11.5, color=NAVY)
    return p

def action_box_header(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.2)
    set_para_shading(p, 'F6F1E9')
    add_run(p, 'YOUR ACTION FOR THIS STEP', bold=True, size=9, color=GOLD, font='Inter')
    return p

def action_item(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.left_indent = Inches(0.35)
    set_para_shading(p, 'F6F1E9')
    add_run(p, f'☐  {text}', size=11, color=BODY)
    return p

def compliance(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(4)
    add_run(p, text, italic=True, size=9, color=MUTED)
    return p

def transition_line(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    add_run(p, text, italic=True, size=10.5, color=MUTED)
    return p

def section_head(doc, text, size=13):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(4)
    add_run(p, text, bold=True, size=size, color=NAVY)
    return p

def bullet_item(doc, label, text, indent=0.2):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.left_indent = Inches(indent)
    if label:
        add_run(p, f'{label}  ', bold=True, size=11, color=NAVY)
        add_run(p, text, size=11, color=BODY)
    else:
        add_run(p, f'•  {text}', size=11, color=BODY)
    return p

def avoid_item(doc, label, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.left_indent = Inches(0.2)
    add_run(p, f'{label}\n', bold=True, size=11, color=NAVY)
    add_run(p, text, size=11, color=BODY)
    return p

def set_doc_margins(doc, top=1.0, bottom=0.9, left=1.1, right=1.1):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

# ────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ────────────────────────────────────────────────────────────

doc = Document()
set_doc_margins(doc)

# ── PAGE 1: COVER ────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.space_after = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_shading(p, '0F1B2D')
add_run(p, 'PROSPER IN AMERICA', bold=True, size=32, color=GOLD, font='Montserrat')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(6)
set_para_shading(p2, '0F1B2D')
add_run(p2, 'The Immigrant Financial Starter Guide', size=16, color=CREAM, font='Inter')

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(4)
p3.paragraph_format.space_after = Pt(4)
set_para_shading(p3, '0F1B2D')
add_run(p3, '7 Steps to Building Your Life in America', italic=True, size=13, color=CREAM, font='Inter')

for _ in range(6):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(0)
    set_para_shading(sp, '0F1B2D')

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(0)
p4.paragraph_format.space_after = Pt(80)
set_para_shading(p4, '0F1B2D')
add_run(p4, 'prosperinamerica.com', size=10, color=MUTED, font='Inter')

doc.add_page_break()

# ── PAGE 2: WELCOME LETTER ───────────────────────────────────
heading(doc, 'A Note Before You Start', size=20, space_before=12)
gold_rule(doc)
spacer(doc, 10)

for line in [
    "You didn’t come here to stay in the same place.",
    "",
    "Neither did I.",
    "",
    "When I arrived in America, no one handed me a roadmap. I figured things out slowly — sometimes the hard way — and I spent years learning things that should have taken weeks.",
    "",
    "This guide is what I wish someone had given me in my first year.",
    "",
    "It won’t answer every question. But it will give you the right starting point — and help you stop making the decisions that cost people the most time and money.",
    "",
    "Read it. Use it. And if you want to apply it to your specific situation, I’m available for a free conversation. No pressure. No agenda.",
]:
    if line == "":
        spacer(doc, 4)
    else:
        body(doc, line, space_before=2, space_after=2)

spacer(doc, 10)
p_sig = doc.add_paragraph()
p_sig.paragraph_format.space_before = Pt(8)
p_sig.paragraph_format.space_after = Pt(2)
add_run(p_sig, '— John David', bold=True, size=12, color=NAVY)
p_brand = doc.add_paragraph()
p_brand.paragraph_format.space_before = Pt(0)
p_brand.paragraph_format.space_after = Pt(6)
add_run(p_brand, 'Prosper In America', size=11, color=GOLD)

doc.add_page_break()

# ── PAGE 3: HOW TO USE ──────────────────────────────────────
heading(doc, 'How to Use This Guide', size=20, space_before=12)
gold_rule(doc)
spacer(doc, 8)

body(doc, 'This guide has 7 steps. They follow a sequence — but you don’t have to start at the beginning.')
spacer(doc, 4)
body(doc, 'If you already have a bank account, skip to Step 3. If credit is your biggest gap, start at Step 4. If you’re worried about your family’s protection, go straight to Step 5.')
spacer(doc, 4)
body(doc, 'Find where you are. Start there.')

spacer(doc, 14)
section_head(doc, 'Steps in this guide')

steps_list = [
    ('01', 'Why Most Immigrants Stay Stuck'),
    ('02', 'Get Your Numbers Right'),
    ('03', 'Open the Right Account'),
    ('04', 'Build Your Credit'),
    ('05', 'Protect What You’re Building'),
    ('06', 'Increase Your Income'),
    ('07', 'Build Something That Lasts'),
]
for num, title in steps_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2)
    set_para_shading(p, 'F6F1E9')
    add_run(p, f'{num}  ', bold=True, size=12, color=GOLD, font='Montserrat')
    add_run(p, title, size=12, color=NAVY, font='Inter')

doc.add_page_break()

# ── STEP 1 ──────────────────────────────────────────────────
step_number(doc, 1)
heading(doc, 'Why Most Immigrants Stay Stuck', size=22, space_before=4)
body(doc, 'It’s not a lack of effort. It’s a lack of structure.', italic=True, color=MUTED, space_before=0, space_after=10)
gold_rule(doc)
spacer(doc, 8)

body(doc, 'You work hard. That’s not the problem.')
spacer(doc, 3)
body(doc, 'You didn’t come here to stay in the same place.')
spacer(doc, 3)
body(doc, 'And still — years pass without real progress.')
spacer(doc, 3)
body(doc, 'The money comes in. The money goes out. Nothing builds.')
spacer(doc, 10)

section_head(doc, 'No one gave you the map.')
body(doc, 'The US financial system does not explain itself. It assumes you already know how credit works, which banks to use, what forms to file, what you qualify for, and in what order to do things.')
spacer(doc, 4)
body(doc, 'If you arrived without that knowledge — or without someone who had it — you’ve been navigating blind.')
spacer(doc, 8)

section_head(doc, 'You’ve been making decisions based on incomplete information.')
for line in [
    'What someone at church mentioned.',
    'What you found in a WhatsApp group.',
    'What you assumed was the same as Brazil.',
    'What you were afraid to ask about.',
]:
    p = body(doc, line, indent=0.25, space_before=2, space_after=2)

spacer(doc, 4)
body(doc, 'Some of those decisions have consequences you haven’t fully felt yet.')
spacer(doc, 10)

section_head(doc, 'Small mistakes compound.')
for line in [
    'The credit you didn’t build in year one is costing you more in year three.',
    'The account you opened because it was easy might be the wrong account.',
    'The filing you skipped — that’s a gap in your record.',
]:
    p = body(doc, line, indent=0.25, space_before=2, space_after=2)

spacer(doc, 4)
body(doc, 'None of this is your fault. It’s the result of operating without a system in a system that expects you to already have one.')
spacer(doc, 12)

section_head(doc, 'What changes when you have the right information')
for line in [
    'You stop making decisions from fear or guesswork.',
    'You know which step comes first — and why.',
    'You stop wasting time on the wrong things.',
    'The path gets clear.',
]:
    body(doc, line, space_before=2, space_after=2)

spacer(doc, 10)
pull_quote(doc, 'The gap is almost never effort. It’s almost always information — and knowing in what order to act on it.')
spacer(doc, 12)

action_box_header(doc)
p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Pt(2)
p_sub.paragraph_format.space_after = Pt(4)
p_sub.paragraph_format.left_indent = Inches(0.2)
set_para_shading(p_sub, 'F6F1E9')
add_run(p_sub, 'Identify the one area where you feel most behind:', size=11, color=BODY)

for item in ['Banking and accounts', 'Credit history', 'Taxes and filing', 'Protecting my family', 'Building income']:
    action_item(doc, item)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(4)
p_note.paragraph_format.space_after = Pt(8)
p_note.paragraph_format.left_indent = Inches(0.35)
set_para_shading(p_note, 'F6F1E9')
add_run(p_note, 'Go to that step first if you need to.', italic=True, size=10, color=MUTED)

spacer(doc, 6)
transition_line(doc, 'Once you understand where you stand, the next step is knowing what you actually have access to.')

doc.add_page_break()

# ── STEP 2 ──────────────────────────────────────────────────
step_number(doc, 2)
heading(doc, 'Get Your Numbers Right', size=22, space_before=4)
body(doc, 'ITIN vs. SSN — what each one is and what each one opens.', italic=True, color=MUTED, space_before=0, space_after=10)
gold_rule(doc)
spacer(doc, 8)

body(doc, 'You arrive knowing you need something to start — but not knowing exactly what it is or what it opens.')
spacer(doc, 12)

section_head(doc, 'SSN — Social Security Number')
body(doc, 'Issued by the Social Security Administration to people authorized to work in the US.')
body(doc, 'If you have work authorization — through a visa, green card, or other status — you may be eligible to apply for one.')
body(doc, 'If you don’t have work authorization, you won’t qualify for an SSN. That’s where the ITIN comes in.')
spacer(doc, 8)

section_head(doc, 'ITIN — Individual Taxpayer Identification Number')
body(doc, 'Issued by the IRS. A tax processing number — not proof of immigration status, not work authorization, and not a path to benefits on its own.')
body(doc, 'It exists for people who have a US tax filing requirement but are not eligible for an SSN.')
spacer(doc, 6)

section_head(doc, 'What the ITIN is commonly used for', size=12)
for item in [
    'Opening certain bank accounts',
    'Filing a US tax return',
    'Building a credit history at some institutions',
    'Applying for certain financial products',
]:
    bullet_item(doc, None, item)

spacer(doc, 8)
section_head(doc, 'How the application generally works', size=12)
body(doc, 'You complete IRS Form W-7 and submit it with your federal tax return and identity documents. Certified professionals called Certifying Acceptance Agents (CAAs) can assist. Many accounting firms and nonprofit organizations offer this service.')
body(doc, 'Processing times vary. Check the IRS website for current estimates.')
spacer(doc, 10)

section_head(doc, 'What this means for you')
for line in [
    'If you have an SSN — you’re already ahead. Make sure you know what it does and doesn’t give you access to.',
    'If you have an ITIN — you have a real starting point. Many banks, lenders, and services accept it.',
    'If you have neither — find out whether the ITIN applies to your situation. It’s the starting point for many people without work authorization.',
]:
    body(doc, line, space_before=2, space_after=4)

spacer(doc, 4)
callout_line(doc, "Don’t wait for the SSN to start building. For many people in this situation, the ITIN is what opens the first real doors.")
spacer(doc, 10)

action_box_header(doc)
for item in [
    'Confirm whether you currently have an SSN or ITIN',
    'If neither, look into the ITIN process — an accountant or CAA can walk you through it',
    'Write down what you need the number for (banking, credit, filing)',
]:
    action_item(doc, item)

spacer(doc, 8)
compliance(doc, 'Note: ITIN eligibility, SSN eligibility, and application requirements depend on individual immigration status and circumstances. General information only — consult a qualified tax professional or immigration specialist.')
transition_line(doc, 'Once your number is in place, the next thing that matters is where your money actually lives.')

doc.add_page_break()

# ── STEP 3 ──────────────────────────────────────────────────
step_number(doc, 3)
heading(doc, 'Open the Right Account', size=22, space_before=4)
gold_rule(doc)
spacer(doc, 10)

p_strong = doc.add_paragraph()
p_strong.paragraph_format.space_before = Pt(4)
p_strong.paragraph_format.space_after = Pt(8)
add_run(p_strong, 'If your account is wrong, everything that comes after gets harder.', bold=True, size=14, color=NAVY)

body(doc, 'You can’t build US credit without one.')
body(doc, 'You can’t pay rent, utilities, or insurance reliably without one.')
body(doc, 'And if you opened the wrong one — or skipped this step entirely — it’s costing you more than you probably realize.')
spacer(doc, 10)

section_head(doc, 'Why the wrong account slows everything down')
body(doc, 'Not all bank accounts are equal. Some are designed for people who are already established. They come with monthly fees, minimum balance requirements, and penalties that quietly drain what you’re trying to build.')
spacer(doc, 4)
body(doc, 'For an immigrant starting from zero, an account with a $15 monthly fee and a $500 minimum balance isn’t just inconvenient.')
spacer(doc, 2)
callout_line(doc, "That’s not a fee. That’s a leak.")
spacer(doc, 6)

section_head(doc, 'What to look for')
for label, text in [
    ('Accepts your documents', 'Many banks accept an ITIN, passport, and proof of US address — no SSN required. Verify directly before applying.'),
    ('No monthly maintenance fees', 'Some accounts charge you just for existing. Look for no monthly fee, or one that’s easily waived.'),
    ('No minimum balance requirement', 'If you’re building from zero, you shouldn’t be penalized for a low balance.'),
    ('Online and mobile access', 'Baseline. If an account doesn’t offer it, keep looking.'),
    ('A path to a credit product', 'Some banks offer credit-builder products to existing customers. This matters more than most people realize.'),
]:
    bullet_item(doc, label + ':', text)

spacer(doc, 8)
section_head(doc, 'What to avoid')
for label, text in [
    ('Prepaid debit cards', "Don’t report to credit bureaus. Don’t build a banking history. A workaround, not a foundation."),
    ('Accounts with stacked fees', 'Monthly maintenance + overdraft + ATM fees. Read the fee schedule before opening anything.'),
    ('Waiting until you have an SSN', 'In many cases, an ITIN and passport are enough. Waiting costs time you don’t need to lose.'),
]:
    bullet_item(doc, label + ':', text)

spacer(doc, 10)
section_head(doc, 'What documents are typically requested', size=12)
for item in ['Passport (valid, unexpired)', 'ITIN or SSN (where required)', 'Proof of US address — lease, utility bill, or official letter']:
    bullet_item(doc, None, item)

spacer(doc, 8)
pull_quote(doc, "The right account doesn’t impress anyone. It just works — without taking money from you while it does.")
spacer(doc, 8)

p_warn = doc.add_paragraph()
p_warn.paragraph_format.space_before = Pt(4)
p_warn.paragraph_format.space_after = Pt(8)
add_run(p_warn, 'If this step is wrong, everything that follows gets slower, more expensive, and harder to fix later.', bold=True, size=11.5, color=NAVY)

action_box_header(doc)
for item in [
    'Research 2–3 account options that accept your current documents',
    'Compare fee structures — look for the fee schedule, not just advertised features',
    'Confirm requirements directly with the institution before applying',
    'Open an account that charges you nothing to exist',
]:
    action_item(doc, item)

spacer(doc, 8)
compliance(doc, 'Note: Account eligibility and document requirements vary by institution and state. Verify requirements directly with any bank or credit union before applying.')
transition_line(doc, 'Once your account is set up correctly, the system starts looking at your credit history.')

doc.add_page_break()

# ── STEP 4 ──────────────────────────────────────────────────
step_number(doc, 4)
heading(doc, 'Build Your Credit', size=22, space_before=4)
gold_rule(doc)
spacer(doc, 8)

p_strong4 = doc.add_paragraph()
p_strong4.paragraph_format.space_before = Pt(4)
p_strong4.paragraph_format.space_after = Pt(8)
add_run(p_strong4, 'Your entire financial history in Brazil counts for nothing here. You start at zero.', bold=True, size=14, color=NAVY)

body(doc, 'Many immigrants arrive with no US credit history. Not bad credit. No credit. In the eyes of the financial system, you don’t exist yet.')
spacer(doc, 4)
body(doc, 'This is why people stay stuck longer than they expect.')
spacer(doc, 4)
body(doc, 'You can’t rush credit history. What you can do is avoid the mistakes that force you to start over — and take the right steps early enough that time works for you.')
spacer(doc, 12)

section_head(doc, 'What actually matters — the three factors')
for label, text in [
    ('Payment history:', 'The single most important factor. Pay on time, every time. One missed payment can stay on your report for years.'),
    ('Credit utilization:', 'How much of your available limit you’re using. High utilization signals financial pressure. Keep balances low.'),
    ('Length of history:', 'Older accounts help. Closing your first account to upgrade is often a mistake.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 4)
body(doc, 'Everything else — new inquiries, credit mix — matters less. Focus on these three first.', italic=True, color=MUTED)
spacer(doc, 10)

section_head(doc, 'How immigrants typically start building credit')
for label, text in [
    ('Secured credit cards:', 'A deposit becomes your credit limit. Use it, pay on time. The payment behavior builds the history — not the deposit.'),
    ('Becoming an authorized user:', 'Someone you trust adds you to their card. Their history can appear on your report — one of the faster ways to establish a starting score.'),
    ('Credit-builder loans:', 'Offered by some banks and credit unions. Fixed monthly payments reported to bureaus. Research availability where you bank.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 10)
section_head(doc, 'What to avoid')
for label, text in [
    ('Missing a payment — even once:', 'Set up automatic minimum payments as a safety net. One late payment stays on record for up to seven years.'),
    ('Maxing out a card:', 'High utilization damages your score even if you pay it off. Keep balances well below your limit.'),
    ('Applying for multiple products at once:', 'Each application creates an inquiry. Research first, apply once.'),
    ('Assuming your Brazilian credit history transfers:', "It doesn’t. The US system has no access to it. That’s not a penalty — it’s how the system works."),
]:
    bullet_item(doc, label, text)

spacer(doc, 4)
callout_line(doc, 'Those who figure it out late often look back on years they can’t recover.')
spacer(doc, 10)

pull_quote(doc, "You can’t rush credit history. What you can do is avoid the mistakes that force you to start over.")
spacer(doc, 10)

action_box_header(doc)
for item in [
    'Check existing US credit history at annualcreditreport.com (free)',
    'Research one secured card or credit-builder product that fits your situation',
    'Set up automatic payments on all existing accounts',
    'Identify someone who might add you as an authorized user — and have that conversation',
]:
    action_item(doc, item)

spacer(doc, 8)
compliance(doc, 'Note: Credit scoring and lending decisions depend on many individual factors. General educational information only — not financial advice. Consult a qualified financial professional.')
transition_line(doc, 'As your financial record starts to build, the next question becomes: what happens if everything stops?')

doc.add_page_break()

# ── STEP 5 ──────────────────────────────────────────────────
step_number(doc, 5)
heading(doc, 'Protect What You’re Building', size=22, space_before=4)
body(doc, 'Many immigrants spend years building — and nothing protecting it.', bold=True, size=13, color=NAVY, space_before=0, space_after=10)
gold_rule(doc)
spacer(doc, 8)

body(doc, 'This is the step people most often skip.')
spacer(doc, 4)
body(doc, 'Not because they don’t care about their family. Because they’re focused on surviving right now, and protection feels like something you deal with later.')
spacer(doc, 4)
body(doc, 'That logic is common. It’s also how families end up in crisis.')
spacer(doc, 10)

section_head(doc, 'What the risk actually looks like')

p_scenario = doc.add_paragraph()
p_scenario.paragraph_format.space_before = Pt(6)
p_scenario.paragraph_format.space_after = Pt(4)
p_scenario.paragraph_format.left_indent = Inches(0.2)
p_scenario.paragraph_format.right_indent = Inches(0.2)
set_para_shading(p_scenario, 'F6F1E9')
add_run(p_scenario, 'You’re the one holding the household together. You work. You send money home. You pay rent. You’re the structure.\n\nNow picture that stopping tomorrow.\n\nIf you couldn’t work for six months — what happens? If you died this year — what happens to your spouse, your children, your family back in Brazil?', size=11.5, color=BODY)

spacer(doc, 6)
body(doc, 'Many families in this situation have no answer to that question. Not because they haven’t thought about it — because they haven’t put anything in place.')
spacer(doc, 4)

p_gap = doc.add_paragraph()
p_gap.paragraph_format.space_before = Pt(4)
p_gap.paragraph_format.space_after = Pt(8)
add_run(p_gap, 'That gap is what protection is for. This is where years of work are either protected — or exposed.', bold=True, size=12, color=NAVY)
spacer(doc, 8)

section_head(doc, 'What protection actually means')
for label, text in [
    ('Life insurance:', 'A policy pays a designated amount to your family if you die. It’s a simple structure: monthly premium, and if the worst happens, your family has coverage for final expenses, lost income, and time to stabilize.'),
    ('Health coverage:', 'Medical debt is one of the leading causes of financial collapse for working families. Options to research: Medicaid, CHIP, ACA marketplace plans. Eligibility varies by immigration status, income, and state.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 10)
section_head(doc, 'Life insurance and immigration status')
body(doc, 'Many families in this situation assume they need an SSN to qualify. That assumption costs them years of coverage they could have had.')
spacer(doc, 4)
body(doc, 'Some life insurance policies accept applicants with an ITIN. The underwriting process varies by policy and provider. But the option exists.')
spacer(doc, 4)
callout_line(doc, 'If you’ve been waiting until you have your SSN, find out what you may be eligible for now.')
spacer(doc, 10)

section_head(doc, 'What to avoid')
for label, text in [
    ('Waiting until you feel stable:', 'There is no stable. Stability is built, not arrived at. People who wait for the right time often wait too long.'),
    ('Assuming you don’t qualify:', 'Many people never check. The assumption is wrong more often than they expect.'),
    ('Choosing coverage based only on price:', 'The cheapest option and the right option are not always the same.'),
    ('Leaving this conversation until something forces it:', 'The time to think about protecting your family is not during a crisis. It’s before one.'),
]:
    bullet_item(doc, label, text)

spacer(doc, 10)
pull_quote(doc, "Protection is not what you build after you’re stable. It’s part of what makes stability possible.")
spacer(doc, 8)

p_delay = doc.add_paragraph()
p_delay.paragraph_format.space_before = Pt(4)
p_delay.paragraph_format.space_after = Pt(10)
add_run(p_delay, 'This step is almost always delayed until something forces the decision. By then, the options are fewer.', bold=True, size=11.5, color=NAVY)

action_box_header(doc)
for item in [
    'Ask honestly: if I couldn’t work for 6 months, what would happen to my family?',
    'Research health coverage — start with healthcare.gov or a local navigator',
    'Look into whether your documents allow you to apply for life insurance',
    'Speak with a licensed professional — not to buy anything, but to understand what’s available',
]:
    action_item(doc, item)

spacer(doc, 8)
compliance(doc, 'Note: Insurance eligibility, coverage terms, and health program eligibility vary by circumstances, immigration status, and state. General informational content only — not insurance advice. Consult a licensed insurance or healthcare professional.')

# ── SAVE ─────────────────────────────────────────────────────
out_path = '/Users/miriampalma/AI-OS/projects/prosper-landing/public/assets/prosper_in_america_guide.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
